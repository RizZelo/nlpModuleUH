"""
Resume/CV Parser Module - FIXED VERSION
Extracts and cleans text from PDF, DOCX, and TXT files.
"""

import re
import os
import logging
from typing import Optional

logger = logging.getLogger(__name__)

# Import libraries for different file formats
try:
    import PyPDF2  # For PDF files
except ImportError:
    PyPDF2 = None

try:
    from docx import Document  # For DOCX files
except ImportError:
    Document = None

try:
    import fitz  # PyMuPDF - better PDF extraction
except ImportError:
    fitz = None


def parse_document(file_path: str) -> Optional[str]:
    """
    Parse a resume/CV file and extract clean text content.
    
    Args:
        file_path (str): Path to the document file (.pdf, .docx, or .txt)
    
    Returns:
        str: Cleaned text content from the document
        None: If parsing fails or file type is not supported
    """
    
    logger.info(f"🔍 Starting to parse: {file_path}")
    
    # Check if file exists
    if not os.path.exists(file_path):
        logger.error(f"❌ Error: File not found - {file_path}")
        return None
    
    logger.info(f"✅ File exists")
    
    # Get file extension to determine file type
    _, file_extension = os.path.splitext(file_path)
    file_extension = file_extension.lower()
    logger.info(f"📄 File extension: {file_extension}")
    
    # Check file size
    file_size = os.path.getsize(file_path)
    logger.info(f"📏 File size: {file_size} bytes")
    
    # Extract raw text based on file type
    try:
        if file_extension == '.pdf':
            logger.info("🔄 Attempting PDF extraction...")
            raw_text = extract_text_from_pdf(file_path)
        elif file_extension in ['.docx', '.doc']:
            logger.info("🔄 Attempting DOCX extraction...")
            raw_text = extract_text_from_docx(file_path)
        elif file_extension == '.txt':
            logger.info("🔄 Attempting TXT extraction...")
            raw_text = extract_text_from_txt(file_path)
        else:
            logger.error(f"❌ Error: Unsupported file type - {file_extension}")
            return None
        
        # Check if text extraction was successful
        if raw_text is None:
            logger.error("❌ Error: Failed to extract text from document")
            return None
        
        logger.info(f"✅ Extracted {len(raw_text)} characters")
        logger.debug(f"📝 First 100 chars: {raw_text[:100]}")
        
        # Clean the extracted text
        cleaned_text = clean_text(raw_text)
        logger.info(f"✅ Cleaned text: {len(cleaned_text)} characters")
        
        return cleaned_text
    
    except Exception as e:
        logger.error(f"❌ Error parsing document: {str(e)}", exc_info=True)
        return None


def extract_text_from_pdf(file_path: str) -> Optional[str]:
    """
    Extract text from PDF - tries PyMuPDF first (better), falls back to PyPDF2
    
    Args:
        file_path (str): Path to the PDF file
    
    Returns:
        str: Raw text extracted from PDF
        None: If extraction fails
    """
    text = ""
    
    # Try PyMuPDF first (much more reliable for complex PDFs)
    if fitz is not None:
        try:
            logger.debug("   Using PyMuPDF for extraction...")
            
            with fitz.open(file_path) as doc:
                logger.debug(f"   PDF has {len(doc)} pages")
                for page_num, page in enumerate(doc):
                    page_text = page.get_text()
                    logger.debug(f"   Page {page_num + 1}: extracted {len(page_text)} chars")
                    text += page_text + "\n"
            
            if text.strip():
                logger.debug(f"   ✅ PyMuPDF successfully extracted {len(text)} chars")
                return text
            else:
                logger.debug("   ⚠️  PyMuPDF extracted empty text, trying PyPDF2...")
        
        except Exception as e:
            logger.debug(f"   ⚠️  PyMuPDF error: {str(e)}, trying PyPDF2...")
    else:
        logger.debug("   ⚠️  PyMuPDF not available, using PyPDF2...")
    
    # Fallback to PyPDF2
    if PyPDF2 is None:
        logger.debug("   ❌ PyPDF2 not installed. Install with: pip install PyPDF2")
        return None
    
    try:
        logger.debug("   Using PyPDF2 for extraction...")
        text = ""
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            logger.debug(f"   PDF has {len(pdf_reader.pages)} pages")
            
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text() or ""
                logger.debug(f"   Page {page_num + 1}: extracted {len(page_text)} chars")
                text += page_text + "\n"
        
        if text.strip():
            logger.debug(f"   ✅ PyPDF2 successfully extracted {len(text)} chars")
            return text
        else:
            logger.debug("   ⚠️  PyPDF2 extracted empty text - PDF might be image-based (scanned)")
            return None
    
    except Exception as e:
        logger.debug(f"   ❌ PyPDF2 error: {str(e)}")
        import traceback
        traceback.print_exc()
        return None


def extract_text_from_docx(file_path: str) -> Optional[str]:
    """
    Extract text from DOCX file with better error handling
    
    Args:
        file_path (str): Path to the DOCX file
    
    Returns:
        str: Raw text extracted from DOCX
        None: If extraction fails
    """
    if Document is None:
        logger.debug("   ❌ python-docx not installed. Install with: pip install python-docx")
        return None
    
    try:
        logger.debug("   Using python-docx for extraction...")
        doc = Document(file_path)
        
        text = ""
        
        # Extract text from paragraphs
        for i, paragraph in enumerate(doc.paragraphs):
            text += paragraph.text + "\n"
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        logger.debug(f"   ✅ Extracted {len(text)} characters from DOCX")
        
        if not text.strip():
            logger.debug("   ⚠️  DOCX appears to be empty")
            return None
            
        return text
    
    except Exception as e:
        logger.debug(f"   ❌ DOCX extraction error: {str(e)}")
        import traceback
        traceback.print_exc()
        return None


def extract_text_from_txt(file_path: str) -> Optional[str]:
    """
    Extract text from a TXT file.
    
    Args:
        file_path (str): Path to the TXT file
    
    Returns:
        str: Raw text from TXT file
        None: If extraction fails
    """
    
    try:
        # Try UTF-8 first
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        
        logger.debug(f"   ✅ Read TXT file with UTF-8 encoding")
        return text
    
    except UnicodeDecodeError:
        # Try with a different encoding if UTF-8 fails
        try:
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
            logger.debug(f"   ✅ Read TXT file with latin-1 encoding")
            return text
        except Exception as e:
            logger.debug(f"   ❌ Error reading TXT file: {str(e)}")
            return None
    
    except Exception as e:
        logger.debug(f"   ❌ Error reading TXT file: {str(e)}")
        return None


def clean_text(text: str) -> str:
    """
    Clean extracted text - gentler approach that preserves more content
    
    Args:
        text (str): Raw text to clean
    
    Returns:
        str: Cleaned text
    """
    if not text:
        return ""
    
    # Replace multiple spaces/tabs with single space (preserve newlines)
    text = re.sub(r'[ \t]+', ' ', text)
    
    # Replace 3+ newlines with just 2 (preserve paragraph breaks)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Remove leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    
    # Remove leading and trailing whitespace from entire text
    text = text.strip()
    
    return text


# Test script
if __name__ == "__main__":
    import sys
    
    logger.debug("="*60)
    logger.debug("CV PARSER TEST")
    logger.debug("="*60)
    
    # Check installed libraries
    logger.debug("\n📦 Checking libraries:")
    if PyPDF2:
        logger.debug("   ✅ PyPDF2 installed")
    else:
        logger.debug("   ❌ PyPDF2 NOT installed - pip install PyPDF2")
    
    if Document:
        logger.debug("   ✅ python-docx installed")
    else:
        logger.debug("   ❌ python-docx NOT installed - pip install python-docx")
    
    if fitz:
        logger.debug("   ✅ PyMuPDF installed (recommended)")
    else:
        logger.debug("   ⚠️  PyMuPDF NOT installed (optional) - pip install PyMuPDF")
    
    # Test with file
    if len(sys.argv) > 1:
        test_file = sys.argv[1]
    else:
        test_file = input("\n📁 Enter file path to test: ").strip()
    
    if not os.path.exists(test_file):
        logger.debug(f"\n❌ File not found: {test_file}")
        sys.exit(1)
    
    logger.debug(f"\n{'='*60}")
    result = parse_document(test_file)
    logger.debug(f"{'='*60}")
    
    if result:
        logger.debug(f"\n✅ SUCCESS!")
        logger.debug(f"📏 Total length: {len(result)} characters")
        logger.debug(f"📏 Word count: {len(result.split())}")
        logger.debug(f"\n📝 First 300 characters:\n")
        print(result[:300])
        logger.debug("\n...")
        if len(result) > 500:
            logger.debug(f"\n📝 Last 200 characters:\n")
            print(result[-200:])
    else:
        logger.debug("\n❌ FAILED to parse document")
        logger.debug("\nTroubleshooting:")
        logger.debug("1. Make sure libraries are installed:")
        logger.debug("   pip install PyPDF2 python-docx PyMuPDF")
        logger.debug("2. Check if the file is corrupted")
        logger.debug("3. For scanned PDFs, you need OCR (tesseract)")